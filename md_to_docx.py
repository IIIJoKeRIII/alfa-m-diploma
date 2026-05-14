# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
import shutil

base = os.path.dirname(os.path.abspath(__file__))
path_md = os.path.join(base, "Курсовая_новостной_портал_React.md")
out_docx = os.path.join(base, "Курсовая_новостной_портал_React.docx")
out_docx_alt = os.path.join(base, "Coursework_News_Portal_React.docx")

with open(path_md, "r", encoding="utf-8") as f:
    lines = f.read().splitlines()

doc = Document()
section = doc.sections[0]
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.left_margin = Mm(25)
section.right_margin = Mm(15)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(14)


def add_para(text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Mm(12.5)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text = text.replace("**", "")
    run = p.add_run(text.strip())
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_heading(text, level):
    h = doc.add_heading(text.strip(), level=min(level, 3))
    for r in h.runs:
        r.font.name = "Times New Roman"


i = 0
while i < len(lines):
    line = lines[i]
    if line.strip() == "---":
        i += 1
        continue
    if line.startswith("# "):
        add_heading(line[2:].strip(), 0)
        i += 1
        continue
    if line.startswith("## "):
        add_heading(line[3:].strip(), 1)
        i += 1
        continue
    if line.startswith("### "):
        add_heading(line[4:].strip(), 2)
        i += 1
        continue
    if "|" in line and line.strip().startswith("|"):
        table_rows = []
        while i < len(lines) and "|" in lines[i]:
            row = [c.strip() for c in lines[i].split("|")]
            row = [c for c in row if c != ""]
            if row and not all(set(c) <= set("- ") for c in row):
                table_rows.append(row)
            i += 1
        if len(table_rows) >= 1:
            ncols = max(len(r) for r in table_rows)
            tbl = doc.add_table(rows=len(table_rows), cols=ncols)
            tbl.style = "Table Grid"
            for ri, row in enumerate(table_rows):
                for ci in range(ncols):
                    tbl.rows[ri].cells[ci].text = row[ci] if ci < len(row) else ""
        continue
    if line.strip() == "":
        i += 1
        continue
    para_lines = []
    while (
        i < len(lines)
        and lines[i].strip() != ""
        and not lines[i].startswith("#")
        and not (lines[i].strip().startswith("|") and "|" in lines[i])
    ):
        para_lines.append(lines[i])
        i += 1
    text = " ".join(para_lines)
    add_para(text)

doc.save(out_docx)
shutil.copy2(out_docx, out_docx_alt)
print("OK", out_docx, os.path.getsize(out_docx))
print("ALT", out_docx_alt)
